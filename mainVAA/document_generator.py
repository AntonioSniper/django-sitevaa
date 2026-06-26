from docx import Document as DocxDocument
from django.http import HttpResponse
from django.conf import settings
import os
from urllib.parse import quote
from .models import WordTemplate

# Папка, где лежат старые Word-шаблоны: SiteVAA/templates/documents/
TEMPLATES_DIR = os.path.join(settings.BASE_DIR, 'templates', 'documents')


def _value(data, field_name):
    return str(getattr(data, field_name, '') or '')


def _find_template(possible_names, contains_words, template_type=None):
    """
    Ищет шаблон.
    1) Сначала берет файл, загруженный через Django-админку.
    2) Если в админке файла нет, ищет старые шаблоны в папке templates/documents.
    """
    if template_type:
        uploaded = WordTemplate.objects.filter(template_type=template_type).first()
        if uploaded and uploaded.file and os.path.exists(uploaded.file.path):
            return uploaded.file.path

    for filename in possible_names:
        path = os.path.join(TEMPLATES_DIR, filename)
        if os.path.exists(path):
            return path

    if os.path.isdir(TEMPLATES_DIR):
        for filename in os.listdir(TEMPLATES_DIR):
            low = filename.lower()
            if low.endswith('.docx') and not low.startswith('~$'):
                if any(word.lower() in low for word in contains_words):
                    return os.path.join(TEMPLATES_DIR, filename)

    existing = []
    if os.path.isdir(TEMPLATES_DIR):
        existing = [name for name in os.listdir(TEMPLATES_DIR) if name.lower().endswith('.docx')]

    raise FileNotFoundError(
        'Не найден Word-шаблон. Загрузите его в админке: /admin/ -> Word-шаблоны. '
        'Или положите файл в папку: {}. Искомые имена: {}. Сейчас в папке: {}'.format(
            TEMPLATES_DIR,
            ', '.join(possible_names),
            ', '.join(existing) if existing else 'нет .docx файлов'
        )
    )


def _replace_in_paragraph(paragraph, replacements):
    """Заменяет плейсхолдеры даже если Word разбил {{ поле }} на разные runs."""
    if not paragraph.runs:
        return

    full_text = ''.join(run.text for run in paragraph.runs)
    new_text = full_text

    for key, value in replacements.items():
        new_text = new_text.replace('{{ ' + key + ' }}', value)
        new_text = new_text.replace('{{' + key + '}}', value)

    if new_text != full_text:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''


def _replace_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, replacements)
            for nested_table in cell.tables:
                _replace_in_table(nested_table, replacements)


def _fill_docx_template(template_path, output_filename, replacements):
    doc = DocxDocument(template_path)

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        _replace_in_table(table, replacements)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        for table in section.header.tables:
            _replace_in_table(table, replacements)
        for paragraph in section.footer.paragraphs:
            _replace_in_paragraph(paragraph, replacements)
        for table in section.footer.tables:
            _replace_in_table(table, replacements)

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = "attachment; filename*=UTF-8''{}".format(quote(output_filename))
    doc.save(response)
    return response


def generate_title_page(data):
    template_path = _find_template(
        ['Титульный.docx', 'Титульный лист.docx', 'Титульник.docx'],
        ['титул'],
        WordTemplate.TITLE,
    )
    replacements = {
        'tip_title': _value(data, 'tip_title'),
        'fio_genitive': _value(data, 'fio_genitive'),
        'module': _value(data, 'module'),
        'specialization': _value(data, 'specialization'),
        'kurs': _value(data, 'kurs'),
        'group': _value(data, 'group'),
        'date_begin': _value(data, 'date_begin'),
        'date_finish': _value(data, 'date_finish'),
        'head1': _value(data, 'head1'),
        'head2': _value(data, 'head2'),
        'ruc_pract': _value(data, 'ruc_pract'),
        'year': _value(data, 'year'),
    }
    return _fill_docx_template(template_path, 'titulny_list.docx', replacements)


def generate_assignment(data):
    template_path = _find_template(
        ['Задание.docx', 'Индивидуальное задание.docx'],
        ['задани'],
        WordTemplate.ASSIGNMENT,
    )
    replacements = {
        'tip': _value(data, 'tip'),
        'familia': _value(data, 'familia'),
        'name': _value(data, 'name'),
        'otchestvo': _value(data, 'otchestvo'),
        'specialization': _value(data, 'specialization'),
        'date_begin': _value(data, 'date_begin'),
        'date_finish': _value(data, 'date_finish'),
        'head1': _value(data, 'head1'),
        'year': _value(data, 'year'),
    }
    return _fill_docx_template(template_path, 'zadanie.docx', replacements)


def generate_diary(data):
    template_path = _find_template(
        ['Дневник.docx', 'Дневник_f7lmcKg.docx'],
        ['дневник'],
        WordTemplate.DIARY,
    )
    replacements = {
        'tip': _value(data, 'tip'),
        'kurs': _value(data, 'kurs'),
        'group': _value(data, 'group'),
        'familia': _value(data, 'familia'),
        'name': _value(data, 'name'),
        'otchestvo': _value(data, 'otchestvo'),
        'specialization': _value(data, 'specialization'),
        'head1': _value(data, 'head1'),
        'head2': _value(data, 'head2'),
        'ruc_pract': _value(data, 'ruc_pract'),
        'module_code': _value(data, 'module_code'),
    }
    return _fill_docx_template(template_path, 'dnevnik.docx', replacements)
